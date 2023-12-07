# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')

# See PyCharm help at https://www.jetbrains.com/help/pycharm/



from docx import Document
import sqlite3

def ler_documento(doc_path):
    document = Document(doc_path)
    for paragraph in document.paragraphs:
        texto = paragraph.text
        # Faça algo com o texto, como exibir ou armazenar em sua base de dados

def criar_tabela():
    conexao = sqlite3.connect('seu_banco_de_dados.db')
    cursor = conexao.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS OAEs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            Km TEXT,
            Linha TEXT
            Cidade TEXT
            Estado TEXT
            Bitola TEXT
            Traçado TEXT
            Trilhos TEXT
            Fixação TEXT
            Comprimento TEXT
            Largura TEXT
            Altura TEXT
            
        )
    ''')
    conexao.commit()
    conexao.close()

def inserir_OAE(Km, Linha, Cidade, Estado, Bitola, Traçado, Trilhos, Fixação, Comprimento, Largura, Altura):
    conexao = sqlite3.connect('seu_banco_de_dados.db')
    cursor = conexao.cursor()
    cursor.execute('INSERT INTO OAEs (Km, Linha, Cidade, Estado, Bitola, Traçado, Trilhos, Fixação, Comprimento, Largura, Altura) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (Km, Linha, Cidade, Estado, Bitola, Traçado, Trilhos, Fixação, Comprimento, Largura, Altura))
    conexao.commit()
    conexao.close()

# Solicita ao usuário para fornecer informações
Km_OAE = input("Digite o Km da OAE: ")
Linha_OAE = input("Digite a Linha da OAE: ")
Cidade_OAE = input("Digite a Cidade da OAE: ")
Estado_OAE = input("Digite o Estado da OAE: ")
Bitola_OAE = input("Digite a Bitola da OAE: ")
Traçado_OAE = input("Digite o Traçado da OAE: ")
Trilhos_OAE = input("Digite os Trilhos da OAE: ")
Fixação_OAE = input("Digite a Fixação da OAE: ")
Comprimento_OAE = input("Digite o Comprimento da OAE: ")
Largura_OAE = input("Digite a Largura da OAE: ")
Altura_OAE = input("Digite a Altura da OAE: ")


# Cria a tabela (execute apenas uma vez)
criar_tabela()

# Insere informações da OAE no banco de dados
inserir_OAE(Km_OAE, Linha_OAE, Cidade_OAE, Estado_OAE, Bitola_OAE, Traçado_OAE, Trilhos_OAE, Fixação_OAE, Comprimento_OAE, Largura_OAE, Altura_OAE)

# Substitua 'caminho/do/seu/arquivo.docx' pelo caminho real do seu arquivo .docx
ler_documento(r'C:\Users\Lorena\Documents\PONTES-MH\ProjetoDeMelhoria\arquivoBasePython\RelatórioMetálica.docx')

import re
from docx import Document

def substituir_marcadores_no_docx(template_path, informacoes_usuario, novo_doc_path):
    document = Document(template_path)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            texto = run.text
            for marcador, valor in informacoes_usuario.items():
                # Usar expressão regular para substituir marcador no texto
                texto = re.sub(fr"\{{{marcador}_OAE\}}", valor, texto)

            # Exibir o texto do parágrafo no console antes da substituição
            print(texto)

            # Atualizar o texto no run
            run.text = texto

    # Salvar o novo documento
    document.save(novo_doc_path)

# Exemplo de uso
informacoes_usuario = {
    "Km_OAE": "123",
    "Linha_OAE": "Linha X",
    "Cidade_OAE": "Exemploville",
    "Estado_OAE": "EX",
    "Bitola_OAE": "1,60 m",
    "Traçado_OAE": "Curvo",
    "Trilhos_OAE": "Trilhos de aço",
    "Fixação_OAE": "Parafusos",
    "Comprimento_OAE": "50 metros",
    "Largura_OAE": "3 metros",
    "Altura_OAE": "5 metros"
}

substituir_marcadores_no_docx(r'C:\Users\Lorena\Documents\PONTES-MH\ProjetoDeMelhoria\arquivoBasePython\RelatórioMetálica.docx', informacoes_usuario, "novo_diagnostico_OAE.docx")




# Agora o código para inserção das patologias

import sqlite3
from docx import Document

def obter_descricoes_patologias_do_banco():
    # Conectar ao banco de dados
    conexao = sqlite3.connect('seu_banco_de_dados_patologias.db')
    cursor = conexao.cursor()

    # Consulta para obter as descrições das patologias
    cursor.execute('SELECT patologia, descricao FROM tabela_patologias')
    resultados = cursor.fetchall()

    # Criar um dicionário com as descrições das patologias
    descricoes_patologias = {patologia: descricao for patologia, descricao in resultados}

    # Fechar a conexão com o banco de dados
    conexao.close()

    return descricoes_patologias

def substituir_patologias_no_docx(doc_path, patologias):
    document = Document(doc_path)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            texto = run.text
            for patologia, descricao in patologias.items():
                # Substituir a chave da patologia no texto pela descrição
                texto = texto.replace(patologia, descricao)

            # Atualizar o texto no run
            run.text = texto

    # Para tabelas no documento
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        texto = run.text
                        for patologia, descricao in patologias.items():
                            # Substituir a chave da patologia no texto pela descrição
                            texto = texto.replace(patologia, descricao)

                        # Atualizar o texto no run
                        run.text = texto

    # Salvar o novo documento
    novo_doc_path = "novo_documento_patologias.db"
    document.save(novo_doc_path)

    print(f"Novo documento com patologias gerado em: {novo_doc_path}")

# Obter as descrições das patologias do banco de dados
descricoes_patologias = obter_descricoes_patologias_do_banco()

# Exemplo de uso
patologias = {
    '{InadequacoesContraventamento}': descricoes_patologias.get('InadequacoesContraventamento', ''),
    '{DeformacoesExcessivas}': descricoes_patologias.get('DeformacoesExcessivas', ''),
    '{CorrosaoMedia}': descricoes_patologias.get('CorrosaoMedia', ''),
    '{SujeiraVegetacao}': descricoes_patologias.get('SujeiraVegetacao', ''),
    '{DesgastePinturaCorrosao}': descricoes_patologias.get('DesgastePinturaCorrosao', ''),
}

# Substitua 'caminho/do/seu/arquivo.docx' pelo caminho real do seu arquivo .docx
substituir_patologias_no_docx(r'C:\Users\Lorena\Documents\PONTES-MH\ProjetoDeMelhoria\arquivoBasePython\RelatórioMetálica.docx', patologias)


import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from PIL import Image, ImageTk
import sqlite3
from docx import Document
from docx.shared import Inches

class AplicacaoDiagnosticoOAE:
    def __init__(self, root):
        self.root = root
        self.root.title("Diagnóstico de OAE")

        # Variáveis para armazenar seleções do usuário
        self.patologias_selecionadas = []
        self.imagens_selecionadas = []

        # Criar widgets
        self.criar_widgets()

    def criar_widgets(self):
        # Label para patologias
        label_patologias = tk.Label(self.root, text="Selecione as patologias/não conformidades identificadas:")
        label_patologias.pack()

        # Checkbuttons para patologias
        patologias = ["InadequacoesContraventamento", "DeformacoesExcessivas", "CorrosaoMedia", "SujeiraVegetacao", "DesgastePinturaCorrosao"]
        for patologia in patologias:
            checkbutton = tk.Checkbutton(self.root, text=patologia, command=lambda p=patologia: self.atualizar_patologias(p))
            checkbutton.pack()

        # Botão para adicionar imagens
        botao_adicionar_imagem = tk.Button(self.root, text="Adicionar Imagem", command=self.adicionar_imagem)
        botao_adicionar_imagem.pack()

        # Botão para gerar relatório
        botao_gerar_relatorio = tk.Button(self.root, text="Gerar Relatório", command=self.gerar_relatorio)
        botao_gerar_relatorio.pack()

    def atualizar_patologias(self, patologia):
        if patologia in self.patologias_selecionadas:
            self.patologias_selecionadas.remove(patologia)
        else:
            self.patologias_selecionadas.append(patologia)

    def adicionar_imagem(self):
        imagem_path = filedialog.askopenfilename(title="Selecione uma Imagem", filetypes=[("Imagens", "*.png;*.jpg;*.jpeg")])

        if imagem_path:
            self.imagens_selecionadas.append(imagem_path)
            messagebox.showinfo("Imagem Adicionada", "Imagem adicionada com sucesso!")

    def gerar_relatorio(self):
        if not self.patologias_selecionadas:
            messagebox.showerror("Erro", "Selecione pelo menos uma patologia/não conformidade!")
            return

        # Conectar ao banco de dados para obter descrições
        descricoes_patologias = self.obter_descricoes_patologias_do_banco()

        # Criar dicionário de patologias selecionadas com descrições
        patologias_descricoes = {patologia: descricoes_patologias[patologia] for patologia in self.patologias_selecionadas}

        # Substituir patologias no documento .docx
        self.substituir_patologias_no_docx(patologias_descricoes)

        # Inserir imagens no documento .docx
        self.inserir_imagens_no_docx()

        messagebox.showinfo("Relatório Gerado", "Relatório gerado com sucesso!")

    def obter_descricoes_patologias_do_banco(self):
        # Conectar ao banco de dados
        conexao = sqlite3.connect('seu_banco_de_dados_patologias.db')
        cursor = conexao.cursor()

        # Consulta para obter as descrições das patologias
        cursor.execute('SELECT patologia, descricao FROM tabela_patologias')
        resultados = cursor.fetchall()

        # Criar um dicionário com as descrições das patologias
        descricoes_patologias = {patologia: descricao for patologia, descricao in resultados}

        # Fechar a conexão com o banco de dados
        conexao.close()

        return descricoes_patologias

    def substituir_patologias_no_docx(self, patologias_descricoes):
        document = Document("caminho/do/seu/arquivo.docx")

        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                texto = run.text
                for patologia, descricao in patologias_descricoes.items():
                    # Substituir a chave da patologia no texto pela descrição
                    texto = texto.replace(patologia, descricao)

                # Atualizar o texto no run
                run.text = texto

        # Salvar o novo documento
        novo_doc_path = "novo_documento_patologias.docx"
        document.save(novo_doc_path)

    def inserir_imagens_no_docx(self):
        document = Document("caminho/do/seu/arquivo.docx")

        # Inserir imagens no documento
        for imagem_path in self.imagens_selecionadas:
            document.add_picture(imagem_path, width=Inches(2))  # Ajuste o tamanho conforme necessário

        # Salvar o novo documento
        novo_doc_path = "novo_documento_patologias_e_imagens.docx"
        document.save(novo_doc_path)

if __name__ == "__main__":
    root = tk.Tk()
    app = AplicacaoDiagnosticoOAE(root)
    root.mainloop()

# Este código cria uma interface gráfica básica com checkbuttons para as patologias,
# um botão para adicionar imagens e um botão para gerar o relatório.
# Quando o usuário seleciona patologias e adiciona imagens,
# o código substitui as patologias no documento .docx e insere as imagens no final do documento.



